import re
from docx import Document
import pdfplumber

def extract_title_universal(doc_path):
    """
    Извлекает название документа из .docx или .pdf файла.
    Пытается обработать:
    1. Документы, начинающиеся со слова "ДОГОВОР" (с пробелами или без, в любом регистре).
    2. Обычные документы (например, Счет-договор) — первая строка до номера/даты.
    3. Документы, где название находится в таблице (например, terds.docx), с учётом регистра и структуры таблицы.
    4. Многострочные записи в ячейках.
    5. PDF-документы.
    6. Двухколонные документы с форматом "Русская часть №______\tАнглийская часть #______".
    """
    # Определяем тип файла по расширению
    if doc_path.lower().endswith('.docx'):
        return extract_title_from_docx(doc_path)
    elif doc_path.lower().endswith('.pdf'):
        return extract_title_from_pdf(doc_path)
    else:
        print(f"Неподдерживаемый формат файла: {doc_path}")
        return ""

def extract_title_from_docx(docx_path):
    """
    Извлекает название из .docx файла.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Ошибка при открытии документа: {e}")
        return ""

    # --- 1. Проверка на документ, начинающийся с "ДОГОВОР" (с пробелами или без, в любом регистре) ---
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            # Приводим к нижнему регистру для проверки
            lower_text = text.lower()
            if lower_text.startswith("д о г о в о р") or lower_text.startswith("договор"):
                # --- Логика обработки "ДОГОВОР" ---
                # Обработка слова "д о г о в о р" -> "договор" в оригинальном регистре
                pattern_dogovor_case_insensitive = r'[Дд]\s+[Оо]\s+[Гг]\s+[Оо]\s+[Вв]\s+[Оо]\s+[Рр]'
                first_line_cleaned = re.sub(pattern_dogovor_case_insensitive, 'ДОГОВОР', text)

                # Удаление номера договора из первой строки
                # Универсальный шаблон для номера: № followed by any combination of letters, numbers, spaces, underscores, hyphens, dots, slashes
                pattern_number = r'\s*№\s*[\w\s\-\/\.]+'
                match = re.search(pattern_number, first_line_cleaned)
                if match:
                    title_part1 = first_line_cleaned[:match.start()].strip()
                else:
                    title_part1 = first_line_cleaned

                # Проверка на продолжение названия во втором параграфе
                title_part2 = ""
                if i + 1 < len(doc.paragraphs):
                    next_paragraph = doc.paragraphs[i + 1]
                    next_text = next_paragraph.text.strip()
                    if next_text and not re.match(r'^[гГ]\.\s*|\b[оО][тТ]\b\s*', next_text):
                        title_part2 = next_text

                full_title = title_part1
                if title_part2:
                    full_title += " " + title_part2
                return re.sub(r'\s+', ' ', full_title).strip()

    # --- 2. Поиск в таблицах ---
    # Перебираем все таблицы в документе
    for table in doc.tables:
        # Перебираем строки таблицы
        for row in table.rows:
            # Перебираем ячейки в строке
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # Приводим к нижнему регистру для проверки
                    lower_cell_text = cell_text.lower()
                    # Проверяем, содержит ли ячейка ключевые слова ДОГОВОР или CONTRACT
                    if "contract" in lower_cell_text or "договор" in lower_cell_text:
                        # --- Обработка текста из ячейки ---
                        # Убираем лишние пробелы и символы
                        cleaned_text = re.sub(r'\s+', ' ', cell_text).strip()

                        # Ищем английскую часть (CONTRACT) в этой ячейке
                        english_part = ""
                        # Проверяем, есть ли в этой ячейке слово "CONTRACT"
                        if "contract" in lower_cell_text:
                            # Удаляем номер (если есть)
                            part_clean = re.sub(r'\s*#\s*[\w\s\-\/\.]*', '', cleaned_text)
                            part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', part_clean)
                            english_part = part_clean.strip()

                        # Ищем русскую часть (ДОГОВОР) в этой ячейке
                        russian_part = ""
                        if "договор" in lower_cell_text:
                            # Удаляем номер
                            part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', cleaned_text)
                            russian_part = part_clean.strip()

                        # Если в этой ячейке найдена только одна часть (английская или русская),
                        # проверяем соседнюю ячейку (по горизонтали) на наличие другой части
                        if (english_part and not russian_part) or (russian_part and not english_part):
                            # Проверяем следующую ячейку в той же строке
                            if j + 1 < len(row.cells):
                                next_cell = row.cells[j + 1]
                                next_cell_text = next_cell.text.strip()
                                if next_cell_text:
                                    lower_next_cell_text = next_cell_text.lower()
                                    # Если в текущей ячейке была английская часть, ищем русскую в следующей
                                    if english_part and not russian_part and "договор" in lower_next_cell_text:
                                        # Очищаем текст следующей ячейки
                                        next_cleaned_text = re.sub(r'\s+', ' ', next_cell_text).strip()
                                        next_part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', next_cleaned_text)
                                        russian_part = next_part_clean.strip()
                                    # Если в текущей ячейке была русская часть, ищем английскую в следующей
                                    elif russian_part and not english_part and "contract" in lower_next_cell_text:
                                        # Очищаем текст следующей ячейки
                                        next_cleaned_text = re.sub(r'\s+', ' ', next_cell_text).strip()
                                        next_part_clean = re.sub(r'\s*#\s*[\w\s\-\/\.]*', '', next_cleaned_text)
                                        next_part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', next_part_clean)
                                        english_part = next_part_clean.strip()

                        # Собираем полное название
                        if english_part and russian_part:
                            full_title = f"{english_part} {russian_part}"
                            return re.sub(r'\s+', ' ', full_title).strip()
                        elif english_part:
                            # Если найдена только английская часть и больше ничего не подходит
                            return english_part
                        elif russian_part:
                            # Если найдена только русская часть и больше ничего не подходит
                            return russian_part
                        # Если в этой ячейке и соседней не нашли обе части, ищем дальше в других ячейках


    # --- 3. Если не найдено в таблицах, ищем первый непустой параграф ---
    for paragraph in doc.paragraphs:
        full_text = paragraph.text.strip()
        if full_text:
            # Приводим к нижнему регистру для проверки шаблона (для "от")
            lower_full_text = full_text.lower()
            # Используем регулярное выражение для поиска " от <дата>" или " № <число>"
            # Теперь проверяем "от" в нижнем регистре
            # Универсальный шаблон для номера/даты
            match = re.search(r'\s+(?:от\s+\d{2}\.\d{2}\.\d{4}|№\s+[\w\s\-\/\.]+)', lower_full_text)

            if match:
                # Берем оригинальный текст до найденной позиции
                title = full_text[:match.start()].strip()
                return title
            else:
                # Альтернативная проверка: разбиваем по словам и ищем "№" или "от" (в нижнем регистре)
                words = full_text.split()
                title_words = []
                for k, word in enumerate(words):
                    lower_word = word.lower()
                    if lower_word == '№' or lower_word == 'от':
                        break
                    # Проверяем, не является ли текущее слово числом после "№" или "от"
                    if word.isdigit() and k > 0 and words[k - 1].lower() in ['№', 'от']:
                        break
                    # Проверяем, не является ли это датой (XX.XX.XXXX) после "от"
                    if re.match(r'\d{2}\.\d{2}\.\d{4}', word) and k > 0 and words[k - 1].lower() == 'от':
                        break
                    title_words.append(word)
                
                if title_words:
                    return " ".join(title_words).strip()

    # 4. Если ничего не подошло, возвращаем пустую строку
    return ""

def extract_title_from_pdf(pdf_path):
    """
    Извлекает название из .pdf файла.
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Перебираем страницы
            for page in pdf.pages:
                # Извлекаем текст со страницы
                text = page.extract_text()
                if text:
                    # Разбиваем текст на строки
                    lines = text.splitlines()
                    for line in lines:
                        line_stripped = line.strip()
                        if line_stripped:
                            # Проверяем, содержит ли строка ключевые слова
                            lower_line = line_stripped.lower()
                            if "contract" in lower_line or "договор" in lower_line:
                                # --- Обработка строки из PDF ---

                                # Убираем лишние пробелы
                                cleaned_line = re.sub(r'\s+', ' ', line_stripped).strip()

                                # Проверяем, является ли строка двухколонной (разделена табуляцией или несколькими пробелами)
                                # Разбиваем по табуляции
                                parts_by_tab = [part.strip() for part in cleaned_line.split('\t') if part.strip()]
                                if len(parts_by_tab) >= 2:
                                    # Пытаемся найти английскую и русскую части среди частей, разделённых табуляцией
                                    english_part = ""
                                    russian_part = ""
                                    for part in parts_by_tab:
                                        lower_part = part.lower()
                                        if "contract" in lower_part and "договор" not in lower_part:
                                            # Удаляем номер
                                            part_clean = re.sub(r'\s*#\s*[\w\s\-\/\.]*', '', part)
                                            part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', part_clean)
                                            english_part = part_clean.strip()
                                        elif "договор" in lower_part:
                                            # Удаляем номер
                                            part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', part)
                                            russian_part = part_clean.strip()

                                    if english_part and russian_part:
                                        full_title = f"{english_part} {russian_part}"
                                        return re.sub(r'\s+', ' ', full_title).strip()

                                # Если не двухколонная строка, или двухколонная но не нашли обе части,
                                # попробуем обработать как обычную строку (как в docx)
                                # Ищем английскую часть
                                english_part = ""
                                if "contract" in lower_line:
                                    part_clean = re.sub(r'\s*#\s*[\w\s\-\/\.]*', '', cleaned_line)
                                    part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', part_clean)
                                    english_part = part_clean.strip()

                                # Ищем русскую часть
                                russian_part = ""
                                if "договор" in lower_line:
                                    part_clean = re.sub(r'\s*№\s*[\w\s\-\/\.]*', '', cleaned_line)
                                    russian_part = part_clean.strip()

                                # Если нашли обе части в одной строке
                                if english_part and russian_part:
                                    full_title = f"{english_part} {russian_part}"
                                    return re.sub(r'\s+', ' ', full_title).strip()
                                elif english_part:
                                    return english_part
                                elif russian_part:
                                    return russian_part

    except Exception as e:
        print(f"Ошибка при обработке PDF: {e}")
        return ""

    # Если ничего не нашли в PDF
    return ""

# # --- Пример использования ---
# file_path1 = "test.docx"
# file_path2 = "test2.docx"
# file_path3 = "terds.docx"
# file_path4 = "example.pdf" # Замените на реальный путь к вашему PDF файлу

# print("--- Обработка test.docx ---")
# title1 = extract_title_universal(file_path1)
# print(f"Извлеченное название: '{title1}'")

# print("\n--- Обработка test2.docx ---")
# title2 = extract_title_universal(file_path2)
# print(f"Извлеченное название: '{title2}'")

# print("\n--- Обработка terds.docx ---")
# title3 = extract_title_universal(file_path3)
# print(f"Извлеченное название: '{title3}'")

# # Для PDF
# if file_path4.endswith('.pdf'):
#     print(f"\n--- Обработка {file_path4} ---")
#     title4 = extract_title_universal(file_path4)
#     print(f"Извлеченное название: '{title4}'")